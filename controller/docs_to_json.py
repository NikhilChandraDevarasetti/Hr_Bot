from docx import *
import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
import csv
import io
import copy
import base64

import json
import string

FAQ_FORMAT = {
    "title": None,
    "id": None,
    "html_tag": None,
    "text": [],
    "faq": False,
    "table_id": None,
    "file": None
}


# Main Docx2JSON


class Docx2Json:

    def _iter_block_items(self, parent, filename):
        # if isinstance(parent, Document):    # doctwo
        # if filename.endswith('docx'):
        parent_elm = parent.element.body
        # elif isinstance(parent, _Cell):
        # elif filename.endswith('xlsx'):
        # parent_elm = parent._tc
        # else:
        # raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _read_docx_tables(self, document, tab_id=None, **kwargs):
        def read_docx_tab(tab, **kwargs):
            vf = io.StringIO()
            writer = csv.writer(vf)
            for row in tab.rows:
                writer.writerow(cell.text for cell in row.cells)
            vf.seek(0)
            return pd.read_csv(vf, **kwargs)

        #    doc = Document(filename)
        if tab_id is None:
            return [read_docx_tab(tab, **kwargs) for tab in document.tables]
        else:
            try:
                return read_docx_tab(document.tables[tab_id], **kwargs)
            except IndexError:
                print(
                    'Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                raise

    def _docx_to_df(self, document, filename):

        main_df = pd.DataFrame(columns=['para_text', 'table_id', 'style'])
        table_mod = pd.DataFrame(columns=['string_value', 'table_id'])
        image_df = pd.DataFrame(columns=['image_index', 'image_rID',
                                         'image_filename', 'image_base64_string'])

        table_list = []
        xml_list = []
        tablecounter = 0
        imagecounter = 0
        blockxmlstring = ''

        for block in self._iter_block_items(document, filename):

            if 'text' in str(block):
                isappend = False

                start_format = ''
                _format = 'Non-Bold'
                for run in block.runs:
                    punctuation = string.punctuation
                    REMOVE_PUNCTUATION = str.maketrans("", "", punctuation)
                    keyword = run.text.translate(REMOVE_PUNCTUATION)
                    if keyword.strip():
                        if run.bold and (not start_format):
                            _format = 'Bold'
                            start_format = 'Bold'
                        elif start_format == 'Bold' and not run.bold:
                            _format = 'Non-Bold'
                        else:
                            _type = "normal"
                            if _format != 'Bold':
                                _format = 'Non-Bold'

                    if not start_format:
                        start_format = _format

                # first bold next is non-bold with text-> heading, para sub
                # first non-bold is bold -> para
                style = str(block.style.name)
                appendtxt = str(block.text)
                appendtxt = appendtxt.replace("\n", "")
                appendtxt = appendtxt.replace("\r", "")
                appendtxt = appendtxt.replace("\t", "")
                tabid = 'Novalue'
                paragraph_split = appendtxt.lower().split()
                isappend = True

                for run in block.runs:
                    xmlstr = str(run.element.xml)
                    my_namespaces = dict([node for _, node in ElementTree.iterparse(
                        StringIO(xmlstr), events=['start-ns'])])
                    root = ET.fromstring(xmlstr)
                    # Check if pic is there in the xml of the element. If yes,
                    # then extract the image data
                    if 'pic:pic' in xmlstr:
                        xml_list.append(xmlstr)
                        for pic in root.findall('.//pic:pic', my_namespaces):
                            cNvPr_elem = pic.find(
                                "pic:nvPicPr/pic:cNvPr", my_namespaces)
                            name_attr = cNvPr_elem.get("name")
                            blip_elem = pic.find(
                                "pic:blipFill/a:blip", my_namespaces)
                            embed_attr = blip_elem.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            isappend = True
                            appendtxt = str(
                                'Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                            document_part = document.part
                            image_part = document_part.related_parts[
                                embed_attr]
                            image_base64 = base64.b64encode(image_part._blob)
                            image_base64 = image_base64.decode()
                            dftemp = pd.DataFrame(
                                {'image_index': [imagecounter], 'image_rID': [embed_attr],
                                 'image_filename': [name_attr],
                                 'image_base64_string': [image_base64]})
                            image_df = image_df.append(dftemp, sort=False)
                            style = 'Novalue'
                        imagecounter = imagecounter + 1

            elif 'table' in str(block):
                isappend = True
                style = 'Novalue'
                tabid = tablecounter
                dfs = self._read_docx_tables(
                    document=document, tab_id=tablecounter)
                appendtxt = dfs.to_html()
                dftemp = pd.DataFrame(
                    {'para_text': [appendtxt], 'table_id': [tablecounter], 'style': [style]})
                table_mod = table_mod.append(dftemp, sort=False)
                table_list.append(dfs)
                tablecounter = tablecounter + 1

            if isappend:
                if 'Heading' in style:
                    style = 'Heading'
                dftemp = pd.DataFrame({'para_text': [appendtxt],
                                       'table_id': [tabid],
                                       'style': [style],
                                       'format': [_format]})
                main_df = main_df.append(dftemp, sort=False)
                main_df = main_df.reset_index(drop=True)
        image_df = image_df.reset_index(drop=True)
        return main_df, image_df

    def _df_to_faq(self, main_df, title):
        id = 0
        final_data = []
        curr_parent_data = None
        curr_data = None
        inner_data_list = []
        parent_id = None

        for i in main_df.itertuples():
            if i.para_text.strip():
                if i.style == 'Heading' or i.format == 'Bold':
                    if inner_data_list:
                        if curr_parent_data:
                            print(curr_parent_data)
                            final_data.append(curr_parent_data)

                        final_data = final_data + inner_data_list
                    curr_parent_data = copy.deepcopy(FAQ_FORMAT)
                    curr_parent_data['id'] = id
                    curr_parent_data['text'] = str(i.para_text)
                    curr_parent_data['title'] = title
                    curr_parent_data['html_tag'] = 'h2'
                    parent_id = id
                    id = id + 1
                    inner_data_list = []
                else:
                    curr_data = copy.deepcopy(FAQ_FORMAT)
                    curr_data['id'] = id
                    curr_data['text'] = str(i.para_text)
                    curr_data['title'] = title
                    curr_data['html_tag'] = 'p'
                    if parent_id:
                        curr_data['parents'].append(parent_id)
                        curr_parent_data['childs'].append(id)
                    id = id + 1
                    inner_data_list.append(curr_data)
        else:
            if curr_parent_data:
                final_data.append(curr_parent_data)
            final_data = final_data + inner_data_list

        return final_data

    def _read_docx(self, filename):
        document = Document(filename)
        return document

    def save_json(self, json_filename, json_data):
        with open(json_filename, 'w+') as f:
            json.dump(json_data, f, indent=4)

    def process(self, filename, file_name):
        title = file_name
        document = self._read_docx(filename=io.BytesIO(filename))
        main_df, image_df = self._docx_to_df(document=document, filename=filename)
        faq_json_data = self._df_to_faq(main_df=main_df, title=title)
        return faq_json_data
